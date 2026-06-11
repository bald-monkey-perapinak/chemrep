"""
Text Extractor — извлечение текста из PDF и DOCX файлов для RAG.

Поддерживаемые форматы:
  - PDF: через PyPDF2
  - DOCX: через python-docx
  - TXT: чтение напрямую

Извлечённый текст разбивается на чанки для последующей индексации embeddings.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)

# Максимальная длина чанка для embeddings
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50


def extract_text(data: bytes, filename: str) -> Optional[str]:
    """
    Извлечь текст из файла.
    Возвращает текст или None, если формат не поддерживается.
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    elif lower.endswith(".docx"):
        return _extract_docx(data)
    elif lower.endswith(".txt"):
        return _extract_txt(data)
    elif lower.endswith(".doc"):
        logger.warning("[TextExtractor] Формат .doc не поддерживается, попробуйте конвертировать в .docx")
        return None

    logger.debug("[TextExtractor] Формат %s не поддерживается для извлечения текста", lower)
    return None


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Разбить текст на перекрывающиеся чанки для embeddings.
    """
    if not text or not text.strip():
        return []

    # Очищаем текст от лишних пробелов и переносов
    text = re.sub(r'\s+', ' ', text).strip()

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]

        # Пытаемся разрезать по предложению
        if end < len(text):
            last_period = chunk.rfind('.')
            last_question = chunk.rfind('?')
            last_excl = chunk.rfind('!')
            sentence_end = max(last_period, last_question, last_excl)
            if sentence_end > chunk_size // 3:
                chunk = chunk[:sentence_end + 1]
                end = start + sentence_end + 1

        if chunk.strip():
            chunks.append(chunk.strip())

        start = end - overlap
        if start >= len(text):
            break

    return chunks


def _extract_pdf(data: bytes) -> Optional[str]:
    """Извлечь текст из PDF через PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(data))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        result = "\n".join(texts)
        logger.info("[TextExtractor] PDF: извлечено %d символов из %d страниц", len(result), len(reader.pages))
        return result if result.strip() else None
    except ImportError:
        logger.warning("[TextExtractor] PyPDF2 не установлен. Установите: pip install PyPDF2")
        return None
    except Exception as e:
        logger.error("[TextExtractor] Ошибка извлечения текста из PDF: %s", e)
        return None


def _extract_docx(data: bytes) -> Optional[str]:
    """Извлечь текст из DOCX через python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        result = "\n".join(texts)
        logger.info("[TextExtractor] DOCX: извлечено %d символов из %d абзацев", len(result), len(doc.paragraphs))
        return result if result.strip() else None
    except ImportError:
        logger.warning("[TextExtractor] python-docx не установлен. Установите: pip install python-docx")
        return None
    except Exception as e:
        logger.error("[TextExtractor] Ошибка извлечения текста из DOCX: %s", e)
        return None


def _extract_txt(data: bytes) -> Optional[str]:
    """Прочитать текстовый файл."""
    try:
        text = data.decode("utf-8")
        logger.info("[TextExtractor] TXT: извлечено %d символов", len(text))
        return text if text.strip() else None
    except UnicodeDecodeError:
        try:
            text = data.decode("cp1251")
            logger.info("[TextExtractor] TXT (cp1251): извлечено %d символов", len(text))
            return text if text.strip() else None
        except Exception:
            logger.warning("[TextExtractor] Не удалось декодировать TXT файл")
            return None
